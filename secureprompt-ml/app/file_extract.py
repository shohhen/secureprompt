"""Text extraction for `/v1/scan-file`.

PDFs: **pdfminer.six** primary + **Tesseract OCR** fallback for image/scanned
pages. (The old `pypdf` extractor mangled icon fonts — FontAwesome ligatures
like "map-marker-alt"/"envelope"/"github" leaked as garbage glued onto the real
text, e.g. `/githuberkinov-wtf`.) DOCX via python-docx; everything else is a
best-effort utf-8/latin-1 decode.

Extraction noise (icon-font ligatures, `(cid:N)` markers, control/private-use
glyphs, leading icon symbols) is stripped by `clean_extracted` so the redaction
output isn't corrupted.
"""
from __future__ import annotations

import io
import logging
import os
import re
import unicodedata

logger = logging.getLogger(__name__)

_CID_RE = re.compile(r"\(cid:\d+\)")
_MULTISPACE_RE = re.compile(r"[ \t]{2,}")
# A standalone page-number footer: `1/16`, `3 / 16`, `12/16`. These are the
# per-page markers pdfminer emits on their own line; XLM-R tags them as
# CREDIT_CARD_EXPIRATION_DATE (MM/YY shape), so we drop them before NER.
_PAGE_NUMBER_RE = re.compile(r"^\d{1,4}\s*/\s*\d{1,4}$")
# A line short enough to be a running header/footer/watermark.
_BOILERPLATE_MAX_LEN = 40
# Identical short line repeated on >= this many pages → boilerplate ("Maxfiy").
_BOILERPLATE_MIN_REPEATS = 3
# A leading icon symbol (single non-word, non-space char) followed by space(s),
# e.g. pdfminer renders a FontAwesome icon as "# email@x.com" or "◦ bullet".
_LEADING_ICON_RE = re.compile(r"^\s*[^\w\s][ \t]+")

# A PDF page yielding fewer than this many chars is treated as image-only and
# sent to OCR (scanned résumés, photos of documents). Env-tunable.
_OCR_PAGE_MIN_CHARS = int(os.getenv("ML_OCR_PAGE_MIN_CHARS", "24"))
_OCR_LANGS = os.getenv("ML_OCR_LANGS", "eng+rus+uzb+uzb_cyrl")
_OCR_DPI = int(os.getenv("ML_OCR_DPI", "200"))


def clean_extracted(text: str) -> str:
    """Strip extraction noise, preserving content (emails, phones, handles).

    Removes `(cid:N)` markers; maps control / private-use / surrogate /
    unassigned codepoints (where icon-font glyphs land) to spaces; drops a
    single leading icon symbol per line; collapses runs of spaces and blank
    lines. A leading `+` immediately followed by digits (a phone) survives,
    because the icon rule requires the symbol to be followed by a space.
    """
    if not text:
        return ""
    text = _CID_RE.sub(" ", text)
    out: list[str] = []
    for ch in text:
        if ch in "\t\n":
            out.append(ch)
            continue
        cat = unicodedata.category(ch)
        if cat == "Cf":
            continue  # zero-width / format: REMOVE (spacing would split tokens
            # like jo​hn@x.com — same rule as text_norm.normalize_for_ner)
        if cat in ("Cc", "Co", "Cs", "Cn"):
            out.append(" ")  # standalone control / private-use icon glyphs
            continue
        out.append(ch)
    lines: list[str] = []
    for ln in "".join(out).split("\n"):
        ln = _LEADING_ICON_RE.sub("", ln)
        ln = _MULTISPACE_RE.sub(" ", ln).strip()
        lines.append(ln)
    return re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip()


def strip_boilerplate_lines(text: str) -> str:
    """Drop per-page boilerplate: standalone page-number footers (`N/M`) and
    short lines repeated across many pages (watermarks like `Maxfiy`).

    Both are noise that survives `clean_extracted` and either produce false
    positives (page numbers → CREDIT_CARD_EXPIRATION_DATE) or clutter the
    redacted output. Content lines — including inline fractions inside prose —
    are untouched because only *standalone* `N/M` lines match.
    """
    if not text:
        return ""
    lines = text.split("\n")
    counts: dict[str, int] = {}
    for ln in lines:
        s = ln.strip()
        if s:
            counts[s] = counts.get(s, 0) + 1

    def _is_boilerplate(line: str) -> bool:
        s = line.strip()
        if not s:
            return False
        if _PAGE_NUMBER_RE.match(s):
            return True
        if len(s) <= _BOILERPLATE_MAX_LEN and counts.get(s, 0) >= _BOILERPLATE_MIN_REPEATS:
            return True
        return False

    return "\n".join(ln for ln in lines if not _is_boilerplate(ln))


def _pdf_page_texts(data: bytes) -> list[str]:
    """Per-page text via pdfminer.six (one parse pass)."""
    from pdfminer.high_level import extract_pages
    from pdfminer.layout import LTTextContainer

    pages: list[str] = []
    for layout in extract_pages(io.BytesIO(data)):
        pages.append(
            "".join(el.get_text() for el in layout if isinstance(el, LTTextContainer))
        )
    return pages


def _ocr_sparse_pages(data: bytes, pages: list[str]) -> list[str]:
    """OCR only the pages whose extracted text is near-empty (image/scanned).

    Best-effort: any failure (missing tesseract/poppler binary, unreadable
    page) leaves the pdfminer text in place. Renders lazily — only reached when
    at least one page looks image-only.
    """
    from pdf2image import convert_from_bytes
    import pytesseract

    images = convert_from_bytes(data, dpi=_OCR_DPI)
    if not pages:
        pages = [""] * len(images)
    for i, img in enumerate(images):
        if i >= len(pages) or len(pages[i].strip()) >= _OCR_PAGE_MIN_CHARS:
            continue
        try:
            ocr = pytesseract.image_to_string(img, lang=_OCR_LANGS)
        except Exception as e:  # noqa: BLE001 — never let OCR break extraction
            logger.warning("OCR failed on page %d (%s)", i, e)
            continue
        if len(ocr.strip()) > len(pages[i].strip()):
            pages[i] = ocr
    return pages


def extract_pdf(data: bytes) -> str:
    try:
        pages = _pdf_page_texts(data)
    except Exception as e:  # noqa: BLE001
        logger.warning("pdfminer extraction failed (%s)", e)
        pages = []
    if not pages or any(len(p.strip()) < _OCR_PAGE_MIN_CHARS for p in pages):
        try:
            pages = _ocr_sparse_pages(data, pages)
        except Exception as e:  # noqa: BLE001 — tesseract/poppler not installed
            logger.warning("OCR fallback unavailable (%s)", e)
    return strip_boilerplate_lines(clean_extracted("\n".join(pages)))


def extract_docx(data: bytes) -> str:
    import docx

    doc = docx.Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return clean_extracted("\n".join(parts))


def decode_best_effort(data: bytes) -> str:
    """Extract text from an uploaded file (PDF / DOCX / plaintext)."""
    head = data[:4]
    if head == b"%PDF":
        try:
            text = extract_pdf(data)
            if text.strip():
                return text
        except Exception as e:  # noqa: BLE001
            logger.warning("PDF extraction failed (%s)", e)
    if head[:2] == b"PK":  # zip container → DOCX
        try:
            text = extract_docx(data)
            if text.strip():
                return text
        except Exception as e:  # noqa: BLE001
            logger.warning("DOCX extraction failed (%s)", e)
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1", errors="replace")
