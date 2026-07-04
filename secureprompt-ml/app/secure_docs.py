"""In-document redaction for DOCX / XLSX / TXT — replace PII strings with
`{{Type_N}}` labels so the underlying text is genuinely removed."""
from __future__ import annotations

import io
from typing import Callable

from app.secure_boxes import byte_to_char_map
from app.secure_labels import build_labels, redact_spans


def _container_spans(texts: list[str], dets: list[dict], labels: dict):
    """Map global byte-offset detections (into ``"\\n".join(texts)``) back to each
    container as ``(char_start, char_end, label)`` spans.

    Detection reports UTF-8 byte offsets into the joined text; we convert to char
    offsets and locate the owning container. Redacting by span (not by matching
    the cleaned display value) means noise inside the original span — NBSP,
    zero-width chars, collapsed whitespace — is removed too, closing the
    fail-open gap where a noisy value never matched its cleaned form. A span
    straddling a container boundary is a v1 limitation and is skipped.
    """
    full = "\n".join(texts)
    b2c = byte_to_char_map(full)
    ranges = []
    pos = 0
    for t in texts:
        ranges.append((pos, pos + len(t)))
        pos += len(t) + 1  # trailing "\n" join char
    per: list[list[tuple[int, int, str]]] = [[] for _ in texts]
    for d in dets:
        cs = b2c.get(d["start"])
        ce = b2c.get(d["end"])
        if cs is None or ce is None or cs >= ce:
            continue
        label = labels.get((d["entity_type"].upper(), d["text"]), "")
        if not label:
            continue
        for i, (a, b) in enumerate(ranges):
            if a <= cs and ce <= b:
                per[i].append((cs - a, ce - a, label))
                break
    return per


def _ordered_labels(dets: list[dict]) -> dict:
    """build_labels in first-appearance (byte-offset) order."""
    return build_labels(sorted(dets, key=lambda d: d.get("start", 0)))


def _set_paragraph_text(paragraph, text: str) -> None:
    """Put `text` in the first direct run (keeps its formatting), blank the rest.

    PII inside ``<w:hyperlink>`` runs is NOT part of ``paragraph.runs`` even
    though it IS part of ``paragraph.text`` — so we must blank hyperlink-nested
    runs too, else the original hyperlinked PII survives in the file (a security
    hole for a redaction control). PII that spanned multiple runs is flattened to
    run[0]'s style, and a redacted hyperlink loses its link formatting — both are
    v1 limitations; text removal (the security property) is unaffected.
    """
    for hl in paragraph.hyperlinks:
        for r in hl.runs:
            r.text = ""
    direct = paragraph.runs
    if direct:
        direct[0].text = text
        for r in direct[1:]:
            r.text = ""
    elif text:
        paragraph.add_run(text)


def _iter_cell_paragraphs(cell):
    """Cell body paragraphs plus paragraphs in any NESTED tables (recursive)."""
    for p in cell.paragraphs:
        yield p
    for table in cell.tables:
        for row in table.rows:
            for c in row.cells:
                yield from _iter_cell_paragraphs(c)


def _txbx_paragraphs(element, parent):
    """Paragraphs inside text boxes (``<w:txbxContent>``) under ``element``.

    Text-box text is NOT part of ``doc.paragraphs`` / header ``.paragraphs`` —
    it lives in the drawing/VML tree — so a name in a letterhead text box would
    survive without this. Best-effort: any structural surprise is swallowed so a
    text box never aborts the whole redaction."""
    try:
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph
        for txbx in element.iter(qn("w:txbxContent")):
            for p in txbx.iterfind(qn("w:p")):
                yield Paragraph(p, parent)
    except Exception:  # noqa: BLE001 — text boxes are a best-effort extra surface
        return


def _iter_paragraphs(doc):
    """Every redactable paragraph: body, tables (incl. nested), headers/footers
    (all first/even/default variants), and text boxes. Deduplicated by paragraph
    XML element so linked-section headers (which share one part) aren't processed
    twice — double-processing would splice a label into already-redacted runs."""
    seen: set[int] = set()

    def _fresh(p):
        k = id(p._p)
        if k in seen:
            return False
        seen.add(k)
        return True

    for p in doc.paragraphs:
        if _fresh(p):
            yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in _iter_cell_paragraphs(cell):
                    if _fresh(p):
                        yield p
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.first_page_header, section.first_page_footer,
                   section.even_page_header, section.even_page_footer):
            if hf is None:
                continue
            for p in hf.paragraphs:
                if _fresh(p):
                    yield p
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in _iter_cell_paragraphs(cell):
                            if _fresh(p):
                                yield p
            for p in _txbx_paragraphs(hf._element, hf):
                if _fresh(p):
                    yield p
    for p in _txbx_paragraphs(doc.element.body, doc):
        if _fresh(p):
            yield p


def secure_docx(data: bytes, detect_fn: Callable[[str], list[dict]]):
    import docx
    doc = docx.Document(io.BytesIO(data))
    paras = list(_iter_paragraphs(doc))
    texts = [p.text for p in paras]
    dets = detect_fn("\n".join(texts))
    labels = _ordered_labels(dets)
    # Redact by span (see _container_spans) so noisy PII is removed even when the
    # cleaned detection value differs from the document text. A value straddling a
    # paragraph/cell boundary is a v1 limitation (skipped by _container_spans).
    for p, t, spans in zip(paras, texts, _container_spans(texts, dets, labels)):
        if spans:
            _set_paragraph_text(p, redact_spans(t, spans))
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue(), dets


def secure_xlsx(data: bytes, detect_fn: Callable[[str], list[dict]]):
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data))
    value_cells = [c for ws in wb.worksheets for row in ws.iter_rows()
                   for c in row if isinstance(c.value, str)]
    comment_cells = [c for ws in wb.worksheets for row in ws.iter_rows()
                     for c in row if c.comment is not None and c.comment.text]
    # cell string values first, then comment texts — both are redactable containers.
    # (Sheet names are left intact: renaming them would break formula references.)
    texts = [c.value for c in value_cells] + [c.comment.text for c in comment_cells]
    dets = detect_fn("\n".join(texts))
    labels = _ordered_labels(dets)
    per = _container_spans(texts, dets, labels)
    n = len(value_cells)
    for c, t, spans in zip(value_cells, texts[:n], per[:n]):
        if spans:
            c.value = redact_spans(t, spans)
    for c, t, spans in zip(comment_cells, texts[n:], per[n:]):
        if spans:
            c.comment.text = redact_spans(t, spans)
    out = io.BytesIO()
    wb.save(out)
    return out.getvalue(), dets


def secure_text(data: bytes, detect_fn: Callable[[str], list[dict]]):
    text = data.decode("utf-8", "replace")
    dets = detect_fn(text)
    labels = _ordered_labels(dets)
    spans = _container_spans([text], dets, labels)[0]
    return redact_spans(text, spans).encode("utf-8"), dets
